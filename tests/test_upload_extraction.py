import io
import sys
from pathlib import Path
from tempfile import TemporaryDirectory
import unittest
from types import SimpleNamespace
from unittest.mock import patch
from zipfile import ZIP_DEFLATED, ZipFile

import httpx
from docx import Document
from fastapi import FastAPI
from reportlab.pdfgen import canvas

from horbot.web.api import (
    _build_preview_url,
    _extract_document_content,
    _render_docx_preview_html,
    _render_pptx_preview_html,
    router as api_router,
)


def create_simple_xlsx(path: Path, text: str) -> None:
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
  <Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>
</Types>
"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
</Relationships>
"""
    workbook = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <sheets>
    <sheet name="Sheet1" sheetId="1" r:id="rId1"/>
  </sheets>
</workbook>
"""
    workbook_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>
</Relationships>
"""
    escaped = (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    sheet = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <sheetData>
    <row r="1">
      <c r="A1" t="inlineStr">
        <is><t>{escaped}</t></is>
      </c>
    </row>
  </sheetData>
</worksheet>
"""
    with ZipFile(path, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", rels)
        archive.writestr("xl/workbook.xml", workbook)
        archive.writestr("xl/_rels/workbook.xml.rels", workbook_rels)
        archive.writestr("xl/worksheets/sheet1.xml", sheet)


def create_simple_pptx(path: Path, text: str) -> None:
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/ppt/presentation.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.presentation.main+xml"/>
  <Override PartName="/ppt/slides/slide1.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slide+xml"/>
</Types>
"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="ppt/presentation.xml"/>
</Relationships>
"""
    presentation = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:presentation xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <p:sldIdLst>
    <p:sldId id="256" r:id="rId1"/>
  </p:sldIdLst>
</p:presentation>
"""
    presentation_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide1.xml"/>
</Relationships>
"""
    escaped = (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    slide = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:sld xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
  <p:cSld>
    <p:spTree>
      <p:sp>
        <p:txBody>
          <a:p>
            <a:r><a:t>{escaped}</a:t></a:r>
          </a:p>
        </p:txBody>
      </p:sp>
    </p:spTree>
  </p:cSld>
</p:sld>
"""
    with ZipFile(path, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", rels)
        archive.writestr("ppt/presentation.xml", presentation)
        archive.writestr("ppt/_rels/presentation.xml.rels", presentation_rels)
        archive.writestr("ppt/slides/slide1.xml", slide)


class UploadExtractionTests(unittest.TestCase):
    def test_extract_pdf_text_with_available_backend(self):
        with TemporaryDirectory() as tmpdir:
            pdf_path = Path(tmpdir) / "sample.pdf"
            pdf = canvas.Canvas(str(pdf_path))
            pdf.drawString(72, 720, "PDF extraction smoke line")
            pdf.save()

            extracted = _extract_document_content(pdf_path, "application/pdf")

            self.assertIsNotNone(extracted)
            self.assertIn("PDF extraction smoke line", extracted)

    def test_extract_docx_text(self):
        with TemporaryDirectory() as tmpdir:
            docx_path = Path(tmpdir) / "sample.docx"
            document = Document()
            document.add_paragraph("DOCX extraction smoke line")
            document.save(docx_path)

            extracted = _extract_document_content(
                docx_path,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

            self.assertIsNotNone(extracted)
            self.assertIn("DOCX extraction smoke line", extracted)

    def test_extract_docx_text_without_python_docx_dependency(self):
        with TemporaryDirectory() as tmpdir:
            docx_path = Path(tmpdir) / "sample.docx"
            document = Document()
            document.add_heading("Fallback Heading", level=1)
            document.add_paragraph("Fallback extraction line")
            document.save(docx_path)

            with patch.dict(sys.modules, {"docx": None}):
                extracted = _extract_document_content(
                    docx_path,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

            self.assertIsNotNone(extracted)
            self.assertIn("Fallback Heading", extracted)
            self.assertIn("Fallback extraction line", extracted)

    def test_extract_xlsx_text(self):
        with TemporaryDirectory() as tmpdir:
            xlsx_path = Path(tmpdir) / "sample.xlsx"
            create_simple_xlsx(xlsx_path, "XLSX extraction smoke line")

            extracted = _extract_document_content(
                xlsx_path,
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

            self.assertIsNotNone(extracted)
            self.assertIn("XLSX extraction smoke line", extracted)

    def test_extract_pptx_text(self):
        with TemporaryDirectory() as tmpdir:
            pptx_path = Path(tmpdir) / "sample.pptx"
            create_simple_pptx(pptx_path, "PPTX extraction smoke line")

            extracted = _extract_document_content(
                pptx_path,
                "application/vnd.openxmlformats-officedocument.presentationml.presentation",
            )

            self.assertIsNotNone(extracted)
            self.assertIn("PPTX extraction smoke line", extracted)

    def test_docx_preview_html_contains_document_text(self):
        with TemporaryDirectory() as tmpdir:
            docx_path = Path(tmpdir) / "preview.docx"
            document = Document()
            document.add_heading("Preview Heading", level=1)
            document.add_paragraph("Preview paragraph body")
            document.save(docx_path)

            rendered = _render_docx_preview_html(docx_path, "preview.docx")

            self.assertIn("Preview Heading", rendered)
            self.assertIn("Preview paragraph body", rendered)
            self.assertIn("<html", rendered)

    def test_docx_preview_html_without_python_docx_dependency(self):
        with TemporaryDirectory() as tmpdir:
            docx_path = Path(tmpdir) / "preview.docx"
            document = Document()
            document.add_heading("Fallback Preview Heading", level=1)
            document.add_paragraph("Fallback Preview paragraph body")
            document.save(docx_path)

            with patch.dict(sys.modules, {"docx": None}):
                rendered = _render_docx_preview_html(docx_path, "preview.docx")

            self.assertIn("Fallback Preview Heading", rendered)
            self.assertIn("Fallback Preview paragraph body", rendered)
            self.assertIn("<html", rendered)

    def test_pptx_preview_html_contains_slide_text(self):
        with TemporaryDirectory() as tmpdir:
            pptx_path = Path(tmpdir) / "preview.pptx"
            create_simple_pptx(pptx_path, "Preview slide body")

            rendered = _render_pptx_preview_html(pptx_path, "preview.pptx")

            self.assertIn("Slide 1", rendered)
            self.assertIn("Preview slide body", rendered)
            self.assertIn("<html", rendered)

    def test_build_preview_url_for_inline_office_documents(self):
        self.assertEqual(
            _build_preview_url(
                "abc",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "document",
            ),
            "/api/files/abc/preview",
        )
        self.assertEqual(
            _build_preview_url(
                "xyz",
                "application/vnd.openxmlformats-officedocument.presentationml.presentation",
                "document",
            ),
            "/api/files/xyz/preview",
        )


class UploadApiTests(unittest.IsolatedAsyncioTestCase):
    async def test_upload_document_does_not_call_minimax_bridge(self):
        app = FastAPI()
        app.include_router(api_router, prefix="/api")
        transport = httpx.ASGITransport(app=app)

        with TemporaryDirectory() as tmpdir:
            upload_dir = Path(tmpdir)
            docx_path = upload_dir / "sample.docx"
            document = Document()
            document.add_paragraph("Upload bridge removal smoke line")
            document.save(docx_path)

            with (
                patch("horbot.web.api._get_upload_dir", return_value=upload_dir),
                patch(
                    "horbot.web.api.get_cached_config",
                    return_value=SimpleNamespace(providers={"minimax": {"apiKey": "fake-key"}}),
                ),
            ):
                async with httpx.AsyncClient(transport=transport, base_url="http://testserver") as client:
                    response = await client.post(
                        "/api/upload",
                        files={
                            "files": (
                                "sample.docx",
                                io.BytesIO(docx_path.read_bytes()),
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                        },
                    )

            self.assertEqual(response.status_code, 200)
            payload = response.json()
            self.assertEqual(len(payload), 1)
            self.assertIsNone(payload[0]["minimax_file_id"])


if __name__ == "__main__":
    unittest.main()
