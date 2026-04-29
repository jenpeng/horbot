"""Upload storage, document extraction, and file preview helpers."""

from __future__ import annotations

from datetime import datetime
import hashlib
import html
import json
import mimetypes
import os
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
import time
import urllib.error
import urllib.request
from urllib.parse import unquote, urlparse
import uuid
from typing import Any, Optional

from fastapi import HTTPException
from loguru import logger
from pydantic import BaseModel

from horbot.utils.helpers import ensure_dir, safe_filename
from horbot.web.message_content import clean_message_content

# File upload configuration
ALLOWED_IMAGE_TYPES = {"image/jpeg", "image/png", "image/gif", "image/webp"}
ALLOWED_VIDEO_TYPES = {"video/mp4", "video/quicktime", "video/x-msvideo"}
ALLOWED_AUDIO_TYPES = {"audio/mpeg", "audio/wav", "audio/mp4", "audio/x-m4a"}
ALLOWED_DOC_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "text/plain",
    "text/markdown",
}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB


def _get_upload_dir() -> Path:
    """Get upload directory, create if not exists."""
    from horbot.utils.paths import get_uploads_dir
    return get_uploads_dir()


def _get_file_category(mime_type: str) -> str:
    """Get file category from mime type."""
    if mime_type in ALLOWED_IMAGE_TYPES:
        return "image"
    elif mime_type in ALLOWED_VIDEO_TYPES:
        return "video"
    elif mime_type in ALLOWED_AUDIO_TYPES:
        return "audio"
    elif mime_type in ALLOWED_DOC_TYPES:
        return "document"
    else:
        return "other"


class UploadResponse(BaseModel):
    """Response model for file upload."""
    file_id: str
    filename: str
    original_name: str
    stored_filename: Optional[str] = None
    mime_type: str
    size: int
    category: str
    url: str
    preview_url: Optional[str] = None
    minimax_file_id: Optional[str] = None  # MiniMax file ID for document processing
    extracted_text: Optional[str] = None  # Extracted text content from documents


INLINE_PREVIEW_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.presentationml.presentation",
}
PPTX_PREVIEW_CACHE_TTL_SECONDS = 3 * 24 * 60 * 60
GENERIC_UPLOAD_MIME_TYPES = {
    "",
    "application/octet-stream",
    "binary/octet-stream",
}


def _resolve_upload_mime_type(*, filename: str, declared_mime_type: str | None = None) -> str:
    """Resolve a stable upload MIME type, falling back to the user-visible extension."""
    declared = str(declared_mime_type or "").strip().lower()
    guessed = mimetypes.guess_type(filename)[0] or ""
    if guessed and (declared in GENERIC_UPLOAD_MIME_TYPES or declared not in ALLOWED_DOC_TYPES | ALLOWED_IMAGE_TYPES | ALLOWED_VIDEO_TYPES | ALLOWED_AUDIO_TYPES):
        return guessed
    return declared or guessed or "application/octet-stream"


def _build_preview_url(file_id: str, mime_type: str, category: str) -> Optional[str]:
    if category == "image" or mime_type in INLINE_PREVIEW_MIME_TYPES:
        return f"/api/files/{file_id}/preview"
    return None


def _get_upload_metadata_dir() -> Path:
    return ensure_dir(_get_upload_dir() / ".meta")


def _get_upload_preview_cache_dir() -> Path:
    return ensure_dir(_get_upload_dir() / ".previews")


def cleanup_upload_preview_cache(
    max_age_seconds: int = PPTX_PREVIEW_CACHE_TTL_SECONDS,
    now: float | None = None,
) -> dict[str, int]:
    """Remove stale PPTX preview cache files.

    Slide PNG directories are intentionally cached across preview sessions, but they
    are large. Keep only recently used cache entries.
    """
    root = _get_upload_preview_cache_dir()
    cutoff = (now or time.time()) - max(0, int(max_age_seconds))
    removed_dirs = 0
    removed_files = 0

    for item in list(root.iterdir()):
        try:
            if item.is_dir():
                if item.name.endswith("-slides") and item.stat().st_mtime < cutoff:
                    shutil.rmtree(item, ignore_errors=True)
                    removed_dirs += 1
                elif item.name.startswith("pptx-preview-") and item.stat().st_mtime < cutoff:
                    shutil.rmtree(item, ignore_errors=True)
                    removed_dirs += 1
                continue
            if item.is_file() and item.suffix.lower() in {".pdf", ".tmp"} and item.stat().st_mtime < cutoff:
                item.unlink()
                removed_files += 1
        except FileNotFoundError:
            continue
        except Exception as exc:
            logger.warning("Failed to cleanup upload preview cache item {}: {}", item, exc)

    if removed_dirs or removed_files:
        logger.info(
            "Cleaned stale upload preview cache: dirs={}, files={}, ttl_seconds={}",
            removed_dirs,
            removed_files,
            max_age_seconds,
        )
    return {"removed_dirs": removed_dirs, "removed_files": removed_files}


def _upload_metadata_path(file_id: str) -> Path:
    return _get_upload_metadata_dir() / f"{file_id}.json"


def _resolve_display_filename(name: str | None, fallback_stem: str = "attachment") -> str:
    candidate = str(name or "").strip()
    if not candidate:
        candidate = fallback_stem
    safe_name = safe_filename(Path(candidate).name) or fallback_stem
    stem = Path(safe_name).stem.strip() or fallback_stem
    suffix = Path(safe_name).suffix
    return f"{stem}{suffix}"


def _dedupe_display_filename(name: str, seen_names: set[str]) -> str:
    candidate = name
    stem = Path(name).stem or "attachment"
    suffix = Path(name).suffix
    index = 2
    normalized_candidate = candidate.casefold()
    while normalized_candidate in seen_names:
        candidate = f"{stem} ({index}){suffix}"
        normalized_candidate = candidate.casefold()
        index += 1
    seen_names.add(normalized_candidate)
    return candidate


def _write_upload_metadata(
    *,
    file_id: str,
    stored_filename: str,
    original_name: str,
    mime_type: str,
    size: int,
    category: str,
) -> None:
    metadata_path = _upload_metadata_path(file_id)
    metadata = {
        "file_id": file_id,
        "stored_filename": stored_filename,
        "original_name": original_name,
        "mime_type": mime_type,
        "size": size,
        "category": category,
    }
    metadata_path.write_text(json.dumps(metadata, ensure_ascii=False), encoding="utf-8")


def _read_upload_metadata(file_id: str) -> dict[str, Any] | None:
    metadata_path = _upload_metadata_path(file_id)
    if not metadata_path.is_file():
        return None
    try:
        data = json.loads(metadata_path.read_text(encoding="utf-8"))
        return data if isinstance(data, dict) else None
    except (OSError, json.JSONDecodeError):
        return None


def _resolve_uploaded_file_by_id(file_id: str) -> tuple[Path, dict[str, Any] | None]:
    upload_dir = _get_upload_dir()
    metadata = _read_upload_metadata(file_id)
    stored_filename = str((metadata or {}).get("stored_filename") or "").strip()
    if stored_filename:
        candidate = upload_dir / stored_filename
        if candidate.is_file():
            return candidate, metadata

    matching_files = [
        path for path in upload_dir.glob(f"{file_id}.*")
        if path.is_file()
    ]
    if not matching_files:
        raise HTTPException(status_code=404, detail="File not found")
    return matching_files[0], metadata


def _resolve_storage_filename(file_info: dict[str, Any]) -> str:
    return str(file_info.get("stored_filename") or file_info.get("filename") or "").strip()


def _build_upload_response_for_path(
    stored_path: Path,
    *,
    file_id: str,
    original_name: str,
) -> UploadResponse:
    mime_type = _resolve_upload_mime_type(filename=original_name or stored_path.name)
    category = _get_file_category(mime_type)
    extracted_text = _extract_document_content(stored_path, mime_type) if category == "document" else None
    return UploadResponse(
        file_id=file_id,
        filename=original_name,
        original_name=original_name,
        stored_filename=stored_path.name,
        mime_type=mime_type,
        size=stored_path.stat().st_size,
        category=category,
        url=f"/api/files/{file_id}",
        preview_url=_build_preview_url(file_id, mime_type, category),
        minimax_file_id=None,
        extracted_text=extracted_text,
    )


def _import_local_media_files(media: list[str] | None) -> list[dict[str, Any]]:
    if not media:
        return []

    upload_dir = _get_upload_dir()
    results: list[dict[str, Any]] = []

    for item in media:
        raw_path = str(item or "").strip()
        if not raw_path:
            continue
        source_path = Path(raw_path).expanduser()
        if not source_path.exists() or not source_path.is_file():
            logger.warning("[ChatAPI] message media path not found, skipping: {}", source_path)
            continue

        file_id = str(uuid.uuid4())
        stored_path = upload_dir / f"{file_id}{source_path.suffix}"
        shutil.copy2(source_path, stored_path)
        mime_type = mimetypes.guess_type(source_path.name)[0] or "application/octet-stream"
        category = _get_file_category(mime_type)
        _write_upload_metadata(
            file_id=file_id,
            stored_filename=stored_path.name,
            original_name=source_path.name,
            mime_type=mime_type,
            size=stored_path.stat().st_size,
            category=category,
        )
        response = _build_upload_response_for_path(
            stored_path,
            file_id=file_id,
            original_name=source_path.name,
        )
        results.append(response.model_dump())
        logger.info("[ChatAPI] Imported outbound media {} -> {}", source_path, stored_path.name)

    return results


REMOTE_IMAGE_HOSTS = {
    "image.pollinations.ai",
    "images.pollinations.ai",
}
REMOTE_IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg", ".avif")
REMOTE_IMAGE_CACHE_TIMEOUT_SECONDS = 12
REMOTE_IMAGE_MAX_BYTES = 20 * 1024 * 1024
_STANDALONE_URL_LINE_PATTERN = re.compile(r"^\s*(?:[-*+]\s+|\d+\.\s+)?(https?://\S+)\s*$", re.IGNORECASE)


def _unwrap_url_token(value: str) -> str:
    candidate = (value or "").strip()
    if candidate.startswith("<") and candidate.endswith(">"):
        return candidate[1:-1].strip()
    return candidate


def _is_likely_remote_image_url(value: str) -> bool:
    candidate = _unwrap_url_token(value)
    if not candidate or not re.match(r"^https?://", candidate, re.IGNORECASE):
        return False

    try:
        parsed = urlparse(candidate)
    except Exception:
        return False

    host = (parsed.hostname or "").lower()
    if host in REMOTE_IMAGE_HOSTS:
        return True

    path = (parsed.path or "").lower()
    return any(path.endswith(ext) for ext in REMOTE_IMAGE_EXTENSIONS)


def _guess_remote_image_file_name(url: str, index: int) -> tuple[str, str, str]:
    parsed = urlparse(url)
    raw_name = Path(parsed.path or "").name
    original_extension = Path(raw_name).suffix.lower()
    extension = original_extension
    if extension not in REMOTE_IMAGE_EXTENSIONS:
        extension = ".jpg"

    stem = Path(raw_name).stem if raw_name else ""
    host = (parsed.hostname or "").lower()
    needs_generated_stem = not stem or stem.lower() in {"prompt", "image"} or original_extension not in REMOTE_IMAGE_EXTENSIONS
    if needs_generated_stem:
        if host in REMOTE_IMAGE_HOSTS:
            seed_match = re.search(r"(?:^|&)seed=(\d+)", parsed.query or "", re.IGNORECASE)
            decoded_prompt = unquote(parsed.path or "")
            if "小马" in decoded_prompt or re.search(r"\bpony\b", decoded_prompt, re.IGNORECASE):
                stem = f"pony-theme-{seed_match.group(1)}" if seed_match else f"pony-theme-{index}"
            else:
                stem = f"pollinations-{seed_match.group(1)}" if seed_match else f"pollinations-{index}"
        else:
            stem = f"remote-image-{index}"

    safe_stem = re.sub(r"[^A-Za-z0-9._-]+", "-", stem).strip("-") or f"remote-image-{index}"
    original_name = f"{safe_stem}{extension}"
    filename = original_name
    mime_type = mimetypes.guess_type(original_name)[0] or "image/jpeg"
    return filename, original_name, mime_type


def _remote_image_file_id(url: str) -> str:
    return f"remote-image-{uuid.uuid5(uuid.NAMESPACE_URL, url).hex[:12]}"


def _remote_image_cache_enabled() -> bool:
    # Keep unit tests deterministic and offline.
    return not bool(os.environ.get("PYTEST_CURRENT_TEST"))


def _list_remote_image_cache_paths() -> list[Path]:
    upload_dir = _get_upload_dir()
    return sorted(upload_dir.glob("remote-image-*.*"))


def _get_remote_image_cache_stats() -> dict[str, Any]:
    cache_files = _list_remote_image_cache_paths()
    total_size_bytes = 0
    newest_updated_at: str | None = None

    for path in cache_files:
        stat = path.stat()
        total_size_bytes += stat.st_size
        updated_at = datetime.fromtimestamp(stat.st_mtime).isoformat()
        if newest_updated_at is None or updated_at > newest_updated_at:
            newest_updated_at = updated_at

    return {
        "count": len(cache_files),
        "total_size_bytes": total_size_bytes,
        "newest_updated_at": newest_updated_at,
    }


def _clear_remote_image_cache() -> dict[str, Any]:
    cache_files = _list_remote_image_cache_paths()
    deleted_count = 0
    deleted_size_bytes = 0

    for path in cache_files:
        try:
            deleted_size_bytes += path.stat().st_size
        except OSError:
            pass
        path.unlink(missing_ok=True)
        _upload_metadata_path(path.stem).unlink(missing_ok=True)
        deleted_count += 1

    return {
        "deleted_count": deleted_count,
        "deleted_size_bytes": deleted_size_bytes,
    }


def _cache_remote_image_file(url: str, index: int) -> dict[str, Any] | None:
    if not _remote_image_cache_enabled():
        return None

    file_id = _remote_image_file_id(url)
    upload_dir = _get_upload_dir()
    existing_files = list(upload_dir.glob(f"{file_id}.*"))
    _, guessed_original_name, guessed_mime_type = _guess_remote_image_file_name(url, index)
    if existing_files:
        return _build_upload_response_for_path(
            existing_files[0],
            file_id=file_id,
            original_name=guessed_original_name,
        ).model_dump()

    request = urllib.request.Request(
        url,
        headers={
            "User-Agent": "Horbot/1.0",
            "Accept": "image/*,*/*;q=0.8",
        },
    )

    try:
        with urllib.request.urlopen(request, timeout=REMOTE_IMAGE_CACHE_TIMEOUT_SECONDS) as response:
            content_type = response.headers.get_content_type() or guessed_mime_type
            if not str(content_type).startswith("image/"):
                return None

            extension = mimetypes.guess_extension(content_type) or Path(guessed_original_name).suffix or ".jpg"
            if extension == ".jpe":
                extension = ".jpg"
            if extension.lower() not in REMOTE_IMAGE_EXTENSIONS:
                extension = Path(guessed_original_name).suffix or ".jpg"

            stored_path = upload_dir / f"{file_id}{extension}"
            temp_path = stored_path.with_suffix(f"{stored_path.suffix}.tmp")
            bytes_written = 0
            with open(temp_path, "wb") as handle:
                while True:
                    chunk = response.read(64 * 1024)
                    if not chunk:
                        break
                    bytes_written += len(chunk)
                    if bytes_written > REMOTE_IMAGE_MAX_BYTES:
                        handle.close()
                        temp_path.unlink(missing_ok=True)
                        logger.warning("[ChatAPI] Remote image too large to cache: {}", url)
                        return None
                    handle.write(chunk)
            temp_path.replace(stored_path)
            original_name = f"{Path(guessed_original_name).stem}{extension}"
            _write_upload_metadata(
                file_id=file_id,
                stored_filename=stored_path.name,
                original_name=original_name,
                mime_type=content_type,
                size=bytes_written,
                category="image",
            )
            logger.info("[ChatAPI] Cached remote image {} -> {}", url, stored_path.name)
            return _build_upload_response_for_path(
                stored_path,
                file_id=file_id,
                original_name=original_name,
            ).model_dump()
    except (urllib.error.URLError, TimeoutError, OSError, ValueError) as exc:
        logger.debug("[ChatAPI] Remote image cache skipped for {}: {}", url, exc)
        return None


def _extract_remote_image_urls(content: str) -> list[str]:
    urls: list[str] = []
    seen: set[str] = set()
    for line in (content or "").splitlines():
        match = _STANDALONE_URL_LINE_PATTERN.match(line)
        if not match:
            continue
        candidate = _unwrap_url_token(match.group(1))
        if not _is_likely_remote_image_url(candidate) or candidate in seen:
            continue
        seen.add(candidate)
        urls.append(candidate)
    return urls


def _strip_standalone_remote_image_url_lines(content: str, urls: list[str]) -> str:
    if not content or not urls:
        return content

    url_set = {_unwrap_url_token(url) for url in urls}
    kept_lines: list[str] = []
    for line in content.splitlines():
        match = _STANDALONE_URL_LINE_PATTERN.match(line)
        if match and _unwrap_url_token(match.group(1)) in url_set:
            continue
        kept_lines.append(line)

    normalized = "\n".join(kept_lines)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def _build_remote_image_files(
    urls: list[str],
    existing_files: list[dict[str, Any]] | None = None,
    *,
    cache_remote: bool = True,
) -> list[dict[str, Any]]:
    files: list[dict[str, Any]] = []
    seen_urls = {
        str((item or {}).get("preview_url") or (item or {}).get("url") or "").strip()
        for item in (existing_files or [])
    }
    seen_file_ids = {
        str((item or {}).get("file_id") or "").strip()
        for item in (existing_files or [])
    }
    for index, url in enumerate(urls, start=1):
        remote_file_id = _remote_image_file_id(url)
        if url in seen_urls or remote_file_id in seen_file_ids:
            continue
        if cache_remote:
            cached_file = _cache_remote_image_file(url, index)
            if cached_file:
                files.append(cached_file)
                seen_file_ids.add(remote_file_id)
                continue
        filename, original_name, mime_type = _guess_remote_image_file_name(url, index)
        files.append({
            "file_id": remote_file_id,
            "filename": filename,
            "original_name": original_name,
            "mime_type": mime_type,
            "size": 0,
            "category": "image",
            "url": url,
            "preview_url": url,
        })
    return files


def _normalize_outbound_content_and_files(
    content: str | None,
    media: list[str] | None = None,
) -> tuple[str, list[dict[str, Any]]]:
    cleaned_content = clean_message_content(content or "")
    outbound_files = _import_local_media_files(media)
    remote_image_urls = _extract_remote_image_urls(cleaned_content)
    if remote_image_urls:
        outbound_files.extend(_build_remote_image_files(remote_image_urls, outbound_files))
        cleaned_content = _strip_standalone_remote_image_url_lines(cleaned_content, remote_image_urls)
    return cleaned_content, outbound_files


def _normalize_saved_assistant_content_and_files(
    content: str | None,
    files: list[dict[str, Any]] | None,
) -> tuple[str, list[dict[str, Any]]]:
    cleaned_content = clean_message_content(content or "")
    normalized_files = list(files or [])
    remote_image_urls = _extract_remote_image_urls(cleaned_content)
    if remote_image_urls:
        normalized_files.extend(_build_remote_image_files(
            remote_image_urls,
            normalized_files,
            cache_remote=False,
        ))
        cleaned_content = _strip_standalone_remote_image_url_lines(cleaned_content, remote_image_urls)
    return cleaned_content, normalized_files


def _extract_text_from_pdf(file_path: Path) -> Optional[str]:
    """Extract text content from PDF file."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n\n".join(text_parts) if text_parts else None
    except ImportError:
        logger.warning("pdfplumber not installed, falling back to PyMuPDF for PDF text extraction")
    except Exception as e:
        logger.error(f"PDF text extraction error with pdfplumber: {e}")

    try:
        import fitz

        text_parts = []
        with fitz.open(file_path) as pdf:
            for page in pdf:
                text = page.get_text("text")
                if text:
                    text_parts.append(text)
        return "\n\n".join(text_parts) if text_parts else None
    except Exception as e:
        logger.error(f"PDF text extraction error with PyMuPDF: {e}")
        return None


def _extract_text_from_docx(file_path: Path) -> Optional[str]:
    """Extract text content from DOCX file."""
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text:
                text_parts.append(para.text)
        return "\n\n".join(text_parts) if text_parts else None
    except ImportError:
        logger.info("python-docx not installed, using XML fallback for DOCX text extraction")
        return _extract_text_from_docx_xml(file_path)
    except Exception as e:
        logger.error(f"DOCX text extraction error: {e}")
        return _extract_text_from_docx_xml(file_path)


def _docx_tag_name(value: str) -> str:
    return value.rsplit("}", 1)[-1]


def _docx_attribute_value(attrs: dict[str, str], suffix: str) -> str:
    for key, value in attrs.items():
        if key.rsplit("}", 1)[-1] == suffix:
            return value
    return ""


def _docx_paragraph_text(node: Any) -> str:
    fragments: list[str] = []
    for child in node.iter():
        tag = _docx_tag_name(getattr(child, "tag", ""))
        if tag == "tab":
            fragments.append("\t")
            continue
        if tag == "br":
            fragments.append("\n")
            continue
        if tag != "t":
            continue
        text = child.text or ""
        if text:
            fragments.append(text)
    return "".join(fragments).strip()


def _docx_paragraph_style(node: Any) -> str:
    for child in node:
        if _docx_tag_name(getattr(child, "tag", "")) != "pPr":
            continue
        for prop in child:
            if _docx_tag_name(getattr(prop, "tag", "")) == "pStyle":
                return _docx_attribute_value(getattr(prop, "attrib", {}), "val").strip().lower()
    return ""


def _parse_docx_blocks(file_path: Path) -> list[dict[str, Any]]:
    try:
        import xml.etree.ElementTree as ET
        from zipfile import ZipFile

        with ZipFile(file_path) as archive:
            if "word/document.xml" not in archive.namelist():
                return []
            root = ET.fromstring(archive.read("word/document.xml"))

        body = next(
            (node for node in root.iter() if _docx_tag_name(getattr(node, "tag", "")) == "body"),
            None,
        )
        if body is None:
            return []

        blocks: list[dict[str, Any]] = []
        for child in body:
            tag = _docx_tag_name(getattr(child, "tag", ""))
            if tag == "p":
                text = _docx_paragraph_text(child)
                if text:
                    blocks.append({
                        "type": "paragraph",
                        "text": text,
                        "style": _docx_paragraph_style(child),
                    })
                continue

            if tag != "tbl":
                continue

            rows: list[list[str]] = []
            for row in child:
                if _docx_tag_name(getattr(row, "tag", "")) != "tr":
                    continue
                cells: list[str] = []
                for cell in row:
                    if _docx_tag_name(getattr(cell, "tag", "")) != "tc":
                        continue
                    cell_lines = [
                        _docx_paragraph_text(paragraph)
                        for paragraph in cell
                        if _docx_tag_name(getattr(paragraph, "tag", "")) == "p"
                    ]
                    cell_text = "\n".join(line for line in cell_lines if line).strip()
                    cells.append(cell_text)
                if any(cell for cell in cells):
                    rows.append(cells)
            if rows:
                blocks.append({"type": "table", "rows": rows})

        return blocks
    except Exception as e:
        logger.error(f"DOCX XML parse error: {e}")
        return []


def _extract_text_from_docx_xml(file_path: Path) -> Optional[str]:
    blocks = _parse_docx_blocks(file_path)
    if not blocks:
        return None

    text_parts: list[str] = []
    for block in blocks:
        if block.get("type") == "paragraph":
            text = str(block.get("text") or "").strip()
            if text:
                text_parts.append(text)
            continue

        if block.get("type") != "table":
            continue
        rows = block.get("rows") or []
        for row in rows:
            values = [str(cell or "").strip() for cell in row]
            if any(values):
                text_parts.append("\t".join(values))

    return "\n\n".join(text_parts) if text_parts else None


def _office_archive_sort_key(name: str) -> tuple[int, str]:
    stem = Path(name).stem
    digits = "".join(ch for ch in stem if ch.isdigit())
    return (int(digits) if digits else 0, name)


def _extract_text_from_pptx(file_path: Path) -> Optional[str]:
    """Extract text content from PPTX file using XML parsing."""
    try:
        import xml.etree.ElementTree as ET
        from zipfile import ZipFile

        with ZipFile(file_path) as archive:
            slide_paths = sorted(
                (
                    name
                    for name in archive.namelist()
                    if name.startswith("ppt/slides/slide") and name.endswith(".xml")
                ),
                key=_office_archive_sort_key,
            )

            text_parts: list[str] = []
            for index, slide_path in enumerate(slide_paths, start=1):
                root = ET.fromstring(archive.read(slide_path))
                slide_texts = [
                    node.text.strip()
                    for node in root.iter()
                    if node.tag.rsplit("}", 1)[-1] == "t" and node.text and node.text.strip()
                ]
                if slide_texts:
                    text_parts.append(f"[Slide {index}]")
                    text_parts.append("\n".join(slide_texts))

        return "\n\n".join(text_parts) if text_parts else None
    except Exception as e:
        logger.error(f"PPTX text extraction error: {e}")
        return None


def _extract_text_from_xlsx(file_path: Path) -> Optional[str]:
    """Extract text content from XLSX file using XML parsing."""
    try:
        import xml.etree.ElementTree as ET
        from zipfile import ZipFile

        with ZipFile(file_path) as archive:
            shared_strings: list[str] = []
            if "xl/sharedStrings.xml" in archive.namelist():
                shared_root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
                for node in shared_root.iter():
                    if node.tag.rsplit("}", 1)[-1] == "t" and node.text and node.text.strip():
                        shared_strings.append(node.text.strip())

            sheet_names: list[str] = []
            if "xl/workbook.xml" in archive.namelist():
                workbook_root = ET.fromstring(archive.read("xl/workbook.xml"))
                for node in workbook_root.iter():
                    if node.tag.rsplit("}", 1)[-1] == "sheet":
                        name = (node.attrib.get("name") or "").strip()
                        if name:
                            sheet_names.append(name)

            worksheet_paths = sorted(
                (
                    name
                    for name in archive.namelist()
                    if name.startswith("xl/worksheets/sheet") and name.endswith(".xml")
                ),
                key=_office_archive_sort_key,
            )

            text_parts: list[str] = []
            for index, sheet_path in enumerate(worksheet_paths):
                root = ET.fromstring(archive.read(sheet_path))
                rows: list[str] = []
                for row in root.iter():
                    if row.tag.rsplit("}", 1)[-1] != "row":
                        continue
                    cell_values: list[str] = []
                    for cell in row:
                        if cell.tag.rsplit("}", 1)[-1] != "c":
                            continue
                        cell_type = (cell.attrib.get("t") or "").strip()
                        value = ""
                        if cell_type == "inlineStr":
                            inline_parts = [
                                node.text.strip()
                                for node in cell.iter()
                                if node.tag.rsplit("}", 1)[-1] == "t" and node.text and node.text.strip()
                            ]
                            value = "".join(inline_parts)
                        else:
                            raw = next(
                                (
                                    node.text.strip()
                                    for node in cell
                                    if node.tag.rsplit("}", 1)[-1] == "v" and node.text and node.text.strip()
                                ),
                                "",
                            )
                            if cell_type == "s" and raw.isdigit():
                                shared_index = int(raw)
                                if 0 <= shared_index < len(shared_strings):
                                    value = shared_strings[shared_index]
                            else:
                                value = raw
                        if value:
                            cell_values.append(value)
                    if cell_values:
                        rows.append("\t".join(cell_values))

                if rows:
                    sheet_label = sheet_names[index] if index < len(sheet_names) else Path(sheet_path).stem
                    text_parts.append(f"[{sheet_label}]")
                    text_parts.append("\n".join(rows))

        return "\n\n".join(text_parts) if text_parts else None
    except Exception as e:
        logger.error(f"XLSX text extraction error: {e}")
        return None


def _extract_text_from_txt(file_path: Path) -> Optional[str]:
    """Extract text content from TXT file."""
    try:
        return file_path.read_text(encoding="utf-8")
    except Exception as e:
        logger.error(f"TXT text extraction error: {e}")
        return None


def _extract_document_content(file_path: Path, mime_type: str) -> Optional[str]:
    """Extract text content from document based on mime type."""
    if mime_type == "application/pdf":
        return _extract_text_from_pdf(file_path)
    elif mime_type in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/docx"):
        return _extract_text_from_docx(file_path)
    elif mime_type == "application/vnd.openxmlformats-officedocument.presentationml.presentation":
        return _extract_text_from_pptx(file_path)
    elif mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        return _extract_text_from_xlsx(file_path)
    elif mime_type in ("text/plain", "text/markdown"):
        return _extract_text_from_txt(file_path)
    return None


def _render_preview_shell(title: str, body_html: str) -> str:
    safe_title = html.escape(title)
    return f"""<!doctype html>
<html lang="en">
  <head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>{safe_title}</title>
    <style>
      :root {{
        color-scheme: light;
        --bg: #f8fafc;
        --panel: #ffffff;
        --panel-soft: #eef2ff;
        --line: #dbe4f0;
        --text: #0f172a;
        --muted: #475569;
        --accent: #2563eb;
      }}
      * {{ box-sizing: border-box; }}
      body {{
        margin: 0;
        background: linear-gradient(180deg, #f8fafc 0%, #eef2ff 100%);
        color: var(--text);
        font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
      }}
      .page {{
        width: min(960px, calc(100vw - 32px));
        margin: 0 auto;
        padding: 24px 0 40px;
      }}
      .hero {{
        margin-bottom: 20px;
        padding: 20px 24px;
        border: 1px solid rgba(37, 99, 235, 0.14);
        border-radius: 24px;
        background: linear-gradient(135deg, rgba(37, 99, 235, 0.10), rgba(14, 165, 233, 0.06));
      }}
      .eyebrow {{
        margin: 0 0 8px;
        font-size: 12px;
        font-weight: 700;
        letter-spacing: 0.12em;
        text-transform: uppercase;
        color: var(--accent);
      }}
      .title {{
        margin: 0;
        font-size: 28px;
        line-height: 1.2;
      }}
      .content {{
        display: grid;
        gap: 16px;
      }}
      .doc-paragraph, .doc-list-item {{
        margin: 0;
        padding: 0;
        color: var(--text);
        line-height: 1.75;
        white-space: pre-wrap;
        word-break: break-word;
      }}
      .doc-heading {{
        margin: 0;
        color: #111827;
        font-weight: 700;
        line-height: 1.35;
      }}
      .doc-card, .slide-card, .table-card {{
        border: 1px solid var(--line);
        border-radius: 22px;
        background: rgba(255, 255, 255, 0.92);
        box-shadow: 0 18px 36px rgba(15, 23, 42, 0.06);
      }}
      .doc-card {{
        padding: 24px;
      }}
      .slide-card {{
        overflow: hidden;
      }}
      .slide-header, .table-header {{
        padding: 14px 18px;
        border-bottom: 1px solid var(--line);
        background: var(--panel-soft);
        color: var(--accent);
        font-size: 13px;
        font-weight: 700;
        letter-spacing: 0.08em;
        text-transform: uppercase;
      }}
      .slide-body, .table-body {{
        padding: 18px;
      }}
      .slide-lines {{
        display: grid;
        gap: 10px;
      }}
      .table-scroll {{
        overflow-x: auto;
      }}
      table {{
        width: 100%;
        border-collapse: collapse;
      }}
      th, td {{
        padding: 10px 12px;
        border: 1px solid var(--line);
        text-align: left;
        vertical-align: top;
        color: var(--text);
        font-size: 14px;
        line-height: 1.55;
      }}
      th {{
        background: #f8fafc;
        font-weight: 700;
      }}
      .text-fallback {{
        margin: 0;
        padding: 20px;
        border: 1px solid var(--line);
        border-radius: 22px;
        background: rgba(255, 255, 255, 0.92);
        color: var(--text);
        font-family: ui-monospace, SFMono-Regular, "SF Mono", Menlo, monospace;
        font-size: 13px;
        line-height: 1.7;
        white-space: pre-wrap;
        word-break: break-word;
      }}
      .empty {{
        padding: 28px;
        border: 1px dashed var(--line);
        border-radius: 22px;
        background: rgba(255, 255, 255, 0.7);
        color: var(--muted);
        text-align: center;
      }}
      @media (max-width: 640px) {{
        .page {{
          width: min(100vw - 20px, 960px);
          padding-top: 12px;
        }}
        .hero {{
          padding: 16px 18px;
          border-radius: 20px;
        }}
        .doc-card {{
          padding: 18px;
        }}
      }}
    </style>
  </head>
  <body>
    <main class="page">
      <section class="hero">
        <p class="eyebrow">Inline preview</p>
        <h1 class="title">{safe_title}</h1>
      </section>
      <section class="content">
        {body_html}
      </section>
    </main>
  </body>
</html>
"""


def _render_text_fallback_preview(title: str, content: Optional[str]) -> str:
    if not content or not content.strip():
        body_html = '<div class="empty">No preview content is available for this file yet.</div>'
    else:
        body_html = f'<pre class="text-fallback">{html.escape(content.strip())}</pre>'
    return _render_preview_shell(title, body_html)


def _render_docx_preview_html(file_path: Path, title: str) -> str:
    try:
        from docx import Document

        document = Document(file_path)
        blocks: list[str] = []

        for para in document.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue
            style_name = ((para.style.name if para.style else "") or "").lower()
            if style_name.startswith("heading"):
                level_match = re.search(r"(\d+)", style_name)
                level = int(level_match.group(1)) if level_match else 2
                level = max(1, min(level, 6))
                blocks.append(f'<h{level} class="doc-heading">{html.escape(text)}</h{level}>')
            else:
                blocks.append(f'<p class="doc-paragraph">{html.escape(text)}</p>')

        for table in document.tables:
            rows_html: list[str] = []
            for row in table.rows:
                cells = [html.escape((cell.text or "").strip()) for cell in row.cells]
                if any(cell for cell in cells):
                    rows_html.append(
                        "<tr>" + "".join(f"<td>{cell or '&nbsp;'}</td>" for cell in cells) + "</tr>"
                    )
            if rows_html:
                blocks.append(
                    '<section class="table-card">'
                    '<div class="table-header">Table</div>'
                    '<div class="table-body"><div class="table-scroll"><table>'
                    + "".join(rows_html)
                    + "</table></div></div></section>"
                )

        if not blocks:
            return _render_text_fallback_preview(
                title,
                _extract_text_from_docx(file_path),
            )

        return _render_preview_shell(title, f'<article class="doc-card">{"".join(blocks)}</article>')
    except ImportError:
        logger.info("python-docx not installed, using XML fallback for DOCX preview")
        return _render_docx_preview_html_from_xml(file_path, title)
    except Exception as e:
        logger.error(f"DOCX preview render error: {e}")
        return _render_docx_preview_html_from_xml(file_path, title)


def _render_docx_preview_html_from_xml(file_path: Path, title: str) -> str:
    blocks = _parse_docx_blocks(file_path)
    rendered_blocks: list[str] = []

    for block in blocks:
        if block.get("type") == "paragraph":
            text = str(block.get("text") or "").strip()
            if not text:
                continue
            style_name = str(block.get("style") or "").strip().lower()
            if style_name.startswith("heading"):
                level_match = re.search(r"(\d+)", style_name)
                level = int(level_match.group(1)) if level_match else 2
                level = max(1, min(level, 6))
                rendered_blocks.append(f'<h{level} class="doc-heading">{html.escape(text)}</h{level}>')
            else:
                rendered_blocks.append(f'<p class="doc-paragraph">{html.escape(text)}</p>')
            continue

        if block.get("type") != "table":
            continue
        rows_html: list[str] = []
        for row in block.get("rows") or []:
            cells = [html.escape(str(cell or "").strip()) for cell in row]
            if any(cell for cell in cells):
                rows_html.append(
                    "<tr>" + "".join(f"<td>{cell or '&nbsp;'}</td>" for cell in cells) + "</tr>"
                )
        if rows_html:
            rendered_blocks.append(
                '<section class="table-card">'
                '<div class="table-header">Table</div>'
                '<div class="table-body"><div class="table-scroll"><table>'
                + "".join(rows_html)
                + "</table></div></div></section>"
            )

    if not rendered_blocks:
        return _render_text_fallback_preview(title, _extract_text_from_docx_xml(file_path))

    return _render_preview_shell(title, f'<article class="doc-card">{"".join(rendered_blocks)}</article>')


def _pptx_slide_paths(archive: Any) -> list[str]:
    return sorted(
        (
            name
            for name in archive.namelist()
            if name.startswith("ppt/slides/slide") and name.endswith(".xml")
        ),
        key=_office_archive_sort_key,
    )


def _pptx_slide_count(file_path: Path) -> int:
    try:
        from zipfile import ZipFile

        with ZipFile(file_path) as archive:
            return len(_pptx_slide_paths(archive))
    except Exception as exc:
        logger.warning("Failed to count PPTX slides for {}: {}", file_path, exc)
        return 0


def _find_soffice_command() -> str | None:
    command = shutil.which("soffice") or shutil.which("libreoffice")
    if command:
        return command
    mac_command = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
    if mac_command.is_file():
        return str(mac_command)
    return None


def _pptx_preview_cache_key(file_path: Path) -> str:
    stat = file_path.stat()
    return hashlib.sha256(
        f"{file_path.resolve()}:{stat.st_size}:{stat.st_mtime_ns}".encode("utf-8", errors="ignore")
    ).hexdigest()[:32]


def _pptx_pdf_preview_cache_path(file_path: Path) -> Path:
    digest = _pptx_preview_cache_key(file_path)
    return _get_upload_preview_cache_dir() / f"{digest}.pdf"


def _pptx_slide_image_cache_dir(file_path: Path) -> Path:
    digest = _pptx_preview_cache_key(file_path)
    return ensure_dir(_get_upload_preview_cache_dir() / f"{digest}-slides")


def _pptx_slide_image_paths(file_path: Path) -> list[Path]:
    image_dir = _pptx_slide_image_cache_dir(file_path)
    return sorted(image_dir.glob("slide-*.png"))


def _touch_preview_cache_path(path: Path) -> None:
    try:
        os.utime(path, None)
    except Exception:
        pass


def _export_pptx_pdf_with_libreoffice(file_path: Path, output_path: Path) -> bool:
    command = _find_soffice_command()
    if not command:
        return False

    try:
        with tempfile.TemporaryDirectory(prefix="pptx-preview-", dir=str(_get_upload_preview_cache_dir())) as tmpdir:
            tmp_root = Path(tmpdir)
            profile_dir = tmp_root / "profile"
            profile_dir.mkdir(parents=True, exist_ok=True)
            safe_input = tmp_root / "input.pptx"
            shutil.copy2(file_path, safe_input)
            result = subprocess.run(
                [
                    command,
                    "--headless",
                    "--nologo",
                    "--nodefault",
                    "--nolockcheck",
                    "--norestore",
                    "--nofirststartwizard",
                    f"-env:UserInstallation={profile_dir.as_uri()}",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(tmp_root),
                    str(safe_input),
                ],
                timeout=300,
                capture_output=True,
                text=True,
                check=False,
            )
            if result.returncode != 0:
                logger.warning(
                    "LibreOffice PPTX preview export failed: code={}, stderr={}",
                    result.returncode,
                    (result.stderr or result.stdout or "").strip()[:500],
                )
                return False
            generated = tmp_root / "input.pdf"
            if not generated.is_file() or generated.stat().st_size <= 0:
                logger.warning("LibreOffice PPTX preview export produced no PDF: {}", generated)
                return False
            generated.replace(output_path)
            return True
    except subprocess.TimeoutExpired:
        logger.warning("LibreOffice PPTX preview export timed out for {}", file_path)
        return False
    except Exception as exc:
        logger.warning("LibreOffice PPTX preview export unavailable: {}", exc)
        return False


def _render_pptx_pdf_preview_path(file_path: Path) -> Path | None:
    """Render PPTX to PDF with a headless converter when available."""
    try:
        cache_path = _pptx_pdf_preview_cache_path(file_path)
    except OSError as exc:
        logger.warning("PPTX preview cache path unavailable: {}", exc)
        return None

    if cache_path.is_file() and cache_path.stat().st_size > 0:
        _touch_preview_cache_path(cache_path)
        return cache_path

    tmp_path = cache_path.with_name(f"{cache_path.stem}.tmp.pdf")
    try:
        if tmp_path.exists():
            tmp_path.unlink()
        if not _export_pptx_pdf_with_libreoffice(file_path, tmp_path):
            return None
        tmp_path.replace(cache_path)
        return cache_path
    except Exception as exc:
        logger.warning("PPTX preview PDF export unavailable: {}", exc)
        return None
    finally:
        try:
            if tmp_path.exists():
                tmp_path.unlink()
        except Exception:
            pass


def _cleanup_pptx_pdf_preview(file_path: Path) -> bool:
    """Remove the intermediate LibreOffice PDF while keeping per-slide PNG caches."""
    try:
        cache_path = _pptx_pdf_preview_cache_path(file_path)
        if cache_path.is_file():
            cache_path.unlink()
            return True
    except Exception as exc:
        logger.warning("Failed to clean PPTX preview PDF for {}: {}", file_path, exc)
    return False


def _render_pdf_pages_to_png(pdf_path: Path, output_dir: Path) -> list[Path]:
    try:
        import fitz

        with fitz.open(pdf_path) as pdf:
            rendered: list[Path] = []
            for index, page in enumerate(pdf, start=1):
                width = float(page.rect.width or 1)
                zoom = max(1.0, min(3.0, 1920.0 / width))
                pixmap = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
                output_path = output_dir / f"slide-{index:03d}.png"
                pixmap.save(output_path)
                rendered.append(output_path)
            return rendered
    except Exception as exc:
        logger.warning("Failed to render PPTX PDF preview pages to PNG: {}", exc)
        return []


def _render_pdf_page_to_png(pdf_path: Path, output_path: Path, page: int) -> Path | None:
    try:
        import fitz

        with fitz.open(pdf_path) as pdf:
            if page < 1 or page > len(pdf):
                return None
            pdf_page = pdf[page - 1]
            width = float(pdf_page.rect.width or 1)
            zoom = max(1.0, min(3.0, 1920.0 / width))
            pixmap = pdf_page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
            pixmap.save(output_path)
            return output_path
    except Exception as exc:
        logger.warning("Failed to render PPTX PDF preview page {} to PNG: {}", page, exc)
        return None


def _pptx_visual_preview_plan(file_path: Path) -> tuple[str, int]:
    """Return the visual renderer and the number of pages it can expose."""
    if _pptx_pdf_preview_cache_path(file_path).is_file() or _find_soffice_command():
        return "libreoffice-pdf-images", max(_pptx_slide_count(file_path), 1)
    return "text-fallback", 0


def _render_pptx_slide_image(file_path: Path, page: int) -> tuple[Path | None, str]:
    """Render one PPTX slide image on demand from the LibreOffice PDF preview."""
    if page < 1:
        return None, "text-fallback"

    output_dir = _pptx_slide_image_cache_dir(file_path)
    output_path = output_dir / f"slide-{page:03d}.png"
    if output_path.is_file() and output_path.stat().st_size > 0:
        engine_path = output_dir / "engine.txt"
        engine = engine_path.read_text(encoding="utf-8").strip() if engine_path.is_file() else "cached"
        if engine == "libreoffice-pdf-images":
            _touch_preview_cache_path(output_dir)
            _touch_preview_cache_path(output_path)
            return output_path, engine

    pdf_path = _render_pptx_pdf_preview_path(file_path)
    if pdf_path and pdf_path.is_file():
        rendered = _render_pdf_page_to_png(pdf_path, output_path, page)
        if rendered:
            (output_dir / "engine.txt").write_text("libreoffice-pdf-images", encoding="utf-8")
            _touch_preview_cache_path(output_dir)
            return rendered, "libreoffice-pdf-images"

    return None, "text-fallback"


def _render_pptx_text_fallback_html(file_path: Path, title: str) -> str:
    """Show an honest fallback when no reliable visual PPTX renderer exists."""
    safe_title = html.escape(title)
    slide_count = _pptx_slide_count(file_path)
    extracted_text = (_extract_text_from_pptx(file_path) or "").strip()
    text_block = (
        f'<pre class="text-fallback">{html.escape(extracted_text)}</pre>'
        if extracted_text
        else '<div class="empty">No extracted slide text is available.</div>'
    )
    body_html = f"""
      <section class="doc-card">
        <p class="eyebrow">PowerPoint Preview</p>
        <h2 class="doc-heading">A reliable PPT renderer is not configured</h2>
        <p class="doc-paragraph">
          Horbot uses LibreOffice to convert PPTX files into PDF and then renders each page as a fixed slide image.
          This avoids inaccurate browser-side PPTX layout reconstruction.
        </p>
        <p class="doc-paragraph">
          Install LibreOffice with <code>./horbot.sh install libreoffice</code> to enable high-fidelity preview for {safe_title}.
          {f"Detected slides: {slide_count}." if slide_count else ""}
        </p>
      </section>
      <section class="doc-card">
        <p class="eyebrow">Extracted Text</p>
        {text_block}
      </section>
    """
    return _render_preview_shell(title, body_html)


def _render_pptx_preview_html(file_id: str, file_path: Path, title: str) -> str:
    engine, visual_pages = _pptx_visual_preview_plan(file_path)
    if visual_pages <= 0:
        return _render_pptx_text_fallback_html(file_path, title)

    safe_title = html.escape(title)
    slide_count = _pptx_slide_count(file_path)
    engine_label = {
        "libreoffice-pdf-images": "LibreOffice PDF images",
    }.get(engine, engine)
    note_html = ""
    js_file_id = json.dumps(file_id)
    return f"""<!doctype html>
<html lang="zh-CN">
  <head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>{safe_title}</title>
    <style>
      :root {{
        color-scheme: light;
        --bg: #0b1120;
        --panel: rgba(15, 23, 42, 0.82);
        --line: rgba(148, 163, 184, 0.28);
        --text: #e5eefc;
        --muted: #94a3b8;
        --accent: #38bdf8;
      }}
      * {{ box-sizing: border-box; }}
      html, body {{ margin: 0; min-height: 100%; background: var(--bg); color: var(--text); }}
      body {{
        font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
        background:
          radial-gradient(circle at 18% 14%, rgba(56, 189, 248, 0.22), transparent 34%),
          radial-gradient(circle at 82% 18%, rgba(251, 146, 60, 0.15), transparent 30%),
          linear-gradient(180deg, #0b1120 0%, #111827 100%);
      }}
      .viewer {{
        min-height: 100vh;
        display: grid;
        grid-template-rows: auto minmax(0, 1fr) auto;
        gap: 14px;
        padding: 16px;
      }}
      .toolbar {{
        display: flex;
        align-items: center;
        justify-content: space-between;
        gap: 12px;
        padding: 12px 14px;
        border: 1px solid var(--line);
        border-radius: 20px;
        background: var(--panel);
        backdrop-filter: blur(16px);
      }}
      .title {{ min-width: 0; }}
      .eyebrow {{
        margin: 0 0 4px;
        color: var(--accent);
        font-size: 11px;
        font-weight: 800;
        letter-spacing: 0.14em;
        text-transform: uppercase;
      }}
      h1 {{
        margin: 0;
        max-width: min(72vw, 980px);
        overflow: hidden;
        text-overflow: ellipsis;
        white-space: nowrap;
        font-size: 17px;
        line-height: 1.35;
      }}
      .meta {{
        display: flex;
        flex-wrap: wrap;
        justify-content: flex-end;
        gap: 8px;
        color: var(--muted);
        font-size: 12px;
      }}
      .pill {{
        border: 1px solid var(--line);
        border-radius: 999px;
        padding: 5px 9px;
        background: rgba(15, 23, 42, 0.55);
      }}
      .stage-wrap {{
        display: grid;
        place-items: center;
        min-height: 0;
      }}
      .stage {{
        position: relative;
        width: min(100%, calc((100vh - 154px) * 16 / 9));
        aspect-ratio: 16 / 9;
        border: 1px solid var(--line);
        border-radius: 24px;
        background: #020617;
        box-shadow: 0 28px 70px rgba(0, 0, 0, 0.38);
        overflow: hidden;
      }}
      .slide {{
        position: absolute;
        inset: 0;
        display: grid;
        place-items: center;
        background: #fff;
      }}
      .slide img {{
        width: 100%;
        height: 100%;
        object-fit: contain;
        background: #fff;
      }}
      .slide.loading::after {{
        content: "Loading slide...";
        position: absolute;
        inset: auto 16px 16px auto;
        border-radius: 999px;
        padding: 7px 11px;
        background: rgba(15, 23, 42, 0.78);
        color: white;
        font-size: 12px;
      }}
      .slide.error {{
        color: #0f172a;
        padding: 24px;
        text-align: center;
      }}
      .nav {{
        display: flex;
        align-items: center;
        justify-content: center;
        gap: 12px;
        padding: 10px 0 0;
      }}
      .nav button {{
        min-width: 44px;
        min-height: 36px;
        border: 1px solid var(--line);
        border-radius: 999px;
        background: rgba(15, 23, 42, 0.72);
        color: var(--text);
        cursor: pointer;
      }}
      .nav button:disabled {{ opacity: 0.35; cursor: not-allowed; }}
      .counter {{ color: var(--muted); font-size: 13px; min-width: 78px; text-align: center; }}
      .pager {{
        display: flex;
        align-items: center;
        justify-content: center;
        gap: 10px;
        padding: 4px 2px 2px;
        color: var(--muted);
        font-size: 13px;
      }}
      .pager input {{
        width: 88px;
        border: 1px solid var(--line);
        border-radius: 12px;
        padding: 8px 10px;
        background: rgba(15, 23, 42, 0.72);
        color: var(--text);
        text-align: center;
      }}
      .notice {{
        margin-top: 10px;
        border: 1px solid rgba(251, 191, 36, 0.35);
        border-radius: 16px;
        padding: 10px 12px;
        background: rgba(251, 191, 36, 0.12);
        color: #fde68a;
        font-size: 13px;
        line-height: 1.55;
      }}
      @media (max-width: 700px) {{
        .viewer {{ padding: 10px; gap: 10px; }}
        .toolbar {{ align-items: flex-start; flex-direction: column; }}
        h1 {{ max-width: 100%; white-space: normal; }}
        .stage {{ width: 100%; }}
      }}
    </style>
  </head>
  <body>
    <main class="viewer">
      <header class="toolbar">
        <div class="title">
          <p class="eyebrow">LibreOffice slide preview</p>
          <h1>{safe_title}</h1>
          {note_html}
        </div>
        <div class="meta">
          <span class="pill" id="counter-pill">1 / {visual_pages}</span>
          <span class="pill">{html.escape(engine_label)}</span>
          {f'<span class="pill">PPTX slides: {slide_count}</span>' if slide_count else ''}
        </div>
      </header>
      <section>
        <div class="stage-wrap">
          <div class="stage" id="stage">
            <section class="slide loading" id="slide">
              <img id="slide-image" alt="Slide 1">
            </section>
          </div>
        </div>
        <div class="nav">
          <button type="button" id="prev" aria-label="Previous slide">←</button>
          <span class="counter" id="counter">1 / {visual_pages}</span>
          <button type="button" id="next" aria-label="Next slide">→</button>
        </div>
      </section>
      <footer class="pager">
        <label for="page-input">Go to slide</label>
        <input id="page-input" type="number" min="1" max="{visual_pages}" value="1" inputmode="numeric">
        <span>of {visual_pages}</span>
      </footer>
    </main>
    <script>
      const totalSlides = {visual_pages};
      const slide = document.getElementById('slide');
      const slideImage = document.getElementById('slide-image');
      const counter = document.getElementById('counter');
      const counterPill = document.getElementById('counter-pill');
      const prev = document.getElementById('prev');
      const next = document.getElementById('next');
      const pageInput = document.getElementById('page-input');
      const fileId = {js_file_id};
      const loaded = new Map();
      let current = 1;

      function slideUrl(page) {{
        return `/api/files/${{encodeURIComponent(fileId)}}/preview/slides/${{page}}`;
      }}

      function cleanupPreviewPdf() {{
        const url = `/api/files/${{encodeURIComponent(fileId)}}/preview/cleanup`;
        if (navigator.sendBeacon) {{
          navigator.sendBeacon(url);
          return;
        }}
        fetch(url, {{ method: 'POST', keepalive: true }}).catch(() => {{}});
      }}

      function preload(page) {{
        if (page < 1 || page > totalSlides || loaded.has(page)) return;
        const image = new Image();
        image.src = slideUrl(page);
        loaded.set(page, image);
      }}

      function show(page) {{
        current = Math.max(1, Math.min(page, totalSlides));
        const label = `${{current}} / ${{totalSlides}}`;
        if (counter) counter.textContent = label;
        if (counterPill) counterPill.textContent = label;
        if (pageInput) pageInput.value = String(current);
        if (prev) prev.disabled = current === 1;
        if (next) next.disabled = current === totalSlides;
        if (slide) {{
          slide.classList.add('loading');
          slide.classList.remove('error');
        }}
        if (slideImage) {{
          slideImage.onload = () => slide?.classList.remove('loading');
          slideImage.onerror = () => {{
            slide?.classList.remove('loading');
            slide?.classList.add('error');
            slideImage.removeAttribute('src');
          }};
          slideImage.alt = `Slide ${{current}}`;
          slideImage.src = slideUrl(current);
        }}
        preload(current - 1);
        preload(current + 1);
      }}

      prev?.addEventListener('click', () => show(current - 1));
      next?.addEventListener('click', () => show(current + 1));
      pageInput?.addEventListener('change', () => show(Number(pageInput.value || current)));
      window.addEventListener('pagehide', cleanupPreviewPdf);
      document.addEventListener('keydown', (event) => {{
        if (event.key === 'ArrowRight' || event.key === ' ' || event.key === 'PageDown') {{
          event.preventDefault();
          show(current + 1);
        }}
        if (event.key === 'ArrowLeft' || event.key === 'PageUp') {{
          event.preventDefault();
          show(current - 1);
        }}
        if (event.key === 'Home') show(1);
        if (event.key === 'End') show(totalSlides);
      }});
      show(1);
    </script>
  </body>
</html>
"""
